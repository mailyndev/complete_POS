from fastapi import APIRouter, Depends, HTTPException, Query, Request
from sqlmodel import Session, select
from typing import List, Dict, Any, Optional
from datetime import datetime, date

from ..database.connection import get_session
from ..database.schema import (
    Proveedor, Compra, DetalleCompra, Producto, Inventario,
    MovimientoInventario, HistorialCostos, Usuario, Auditoria, Lote, CuentaPorPagar, get_cr_time
)
from ..utils.security import get_current_user, PermissionChecker
from ..utils.mail_sender import send_email_with_pdf

router = APIRouter(prefix="/purchases", tags=["Compras"])

@router.get("/providers", response_model=List[Proveedor])
def get_providers(
    search: Optional[str] = Query(None, description="Buscar por identificación, nombre, contacto, teléfono, correo o dirección"),
    include_inactive: Optional[bool] = Query(False, description="Incluir proveedores inactivos"),
    current_user: Usuario = Depends(get_current_user),
    session: Session = Depends(get_session)
):
    query = select(Proveedor)
    if not include_inactive:
        query = query.where(Proveedor.activo == True)
        
    if search:
        search_filter = f"%{search}%"
        query = query.where(
            Proveedor.nombre.like(search_filter) | 
            Proveedor.identificacion.like(search_filter) |
            Proveedor.contacto.like(search_filter) |
            Proveedor.telefono.like(search_filter) |
            Proveedor.correo.like(search_filter) |
            Proveedor.direccion.like(search_filter)
        )
    return session.exec(query).all()

@router.post("/providers", status_code=201)
def create_provider(
    payload: Dict[str, Any],
    request: Request,
    current_user: Usuario = Depends(PermissionChecker("purchases:access")),
    session: Session = Depends(get_session)
):
    identificacion = payload.get("identificacion")
    nombre = payload.get("nombre")
    if not identificacion or not nombre:
        raise HTTPException(status_code=400, detail="Identificación y nombre son requeridos.")
        
    dup = session.exec(select(Proveedor).where(Proveedor.identificacion == identificacion)).first()
    if dup:
        raise HTTPException(status_code=400, detail="Ya existe un proveedor con esta identificación.")
        
    prov = Proveedor(
        identificacion=identificacion,
        nombre=nombre,
        contacto=payload.get("contacto", ""),
        telefono=payload.get("telefono", ""),
        correo=payload.get("correo", ""),
        direccion=payload.get("direccion", "")
    )
    session.add(prov)
    session.commit()
    session.refresh(prov)
    
    # Auditoría
    audit = Auditoria(
        usuario_id=current_user.id,
        accion="CREAR_PROVEEDOR",
        modulo="Compras",
        detalles=f"Proveedor creado: {prov.nombre} ({prov.identificacion})",
        ip_address=request.client.host
    )
    session.add(audit)
    session.commit()
    
    return {"message": "Proveedor creado exitosamente", "id": prov.id}

@router.post("", status_code=201)
def record_purchase(
    payload: Dict[str, Any],
    request: Request,
    current_user: Usuario = Depends(PermissionChecker("purchases:access")),
    session: Session = Depends(get_session)
):
    proveedor_id = payload.get("proveedor_id")
    numero_factura = payload.get("numero_factura")
    items = payload.get("items", []) # List[Dict[str, Any]]: producto_id, cantidad, costo_unitario
    estado = payload.get("estado", "pendiente")
    
    if not proveedor_id or not numero_factura or not items:
        raise HTTPException(status_code=400, detail="Proveedor, factura e items son requeridos.")
        
    # Validar factura duplicada
    dup_invoice = session.exec(select(Compra).where(Compra.proveedor_id == proveedor_id, Compra.numero_factura == numero_factura)).first()
    if dup_invoice:
        raise HTTPException(status_code=400, detail="Esta factura ya fue registrada para este proveedor.")
        
    # Validar items
    for item in items:
        qty = float(item.get("cantidad", 0))
        costo = float(item.get("costo_unitario", 0))
        if qty <= 0:
            raise HTTPException(status_code=400, detail="La cantidad de los productos debe ser mayor a cero.")
        if costo < 0:
            raise HTTPException(status_code=400, detail="El costo unitario no puede ser negativo.")
            
    total = sum(float(i["cantidad"]) * float(i["costo_unitario"]) for i in items)
    
    compra = Compra(
        proveedor_id=proveedor_id,
        sucursal_id=current_user.sucursal_id,
        usuario_id=current_user.id,
        numero_factura=numero_factura,
        fecha_compra=get_cr_time(),
        total=total,
        estado=estado
    )
    session.add(compra)
    session.commit()
    session.refresh(compra)
    
    # Registrar Cuenta por Pagar si la compra es a crédito (pendiente)
    if estado == "pendiente":
        from datetime import date, timedelta
        cxp = CuentaPorPagar(
            proveedor_id=proveedor_id,
            compra_id=compra.id,
            monto_total=total,
            saldo_pendiente=total,
            fecha_vencimiento=date.today() + timedelta(days=30),
            estado="pendiente"
        )
        session.add(cxp)
        
    for item in items:
        prod_id = item["producto_id"]
        qty = float(item["cantidad"])
        costo = float(item["costo_unitario"])
        
        detalle = DetalleCompra(
            compra_id=compra.id,
            producto_id=prod_id,
            cantidad=qty,
            costo_unitario=costo
        )
        session.add(detalle)
        
        # Aumentar existencia en el inventario general
        inv = session.exec(select(Inventario).where(
            Inventario.producto_id == prod_id,
            Inventario.sucursal_id == current_user.sucursal_id
        )).first()
        
        if not inv:
            inv = Inventario(sucursal_id=current_user.sucursal_id, producto_id=prod_id, existencia_actual=0.0)
            
        inv.existencia_actual += qty
        session.add(inv)
        session.commit()
        session.refresh(inv)
        
        # Registrar Kárdex (Entrada)
        mov = MovimientoInventario(
            inventario_id=inv.id,
            tipo_movimiento="entrada",
            cantidad=qty,
            motivo=f"Entrada por compra. Factura Proveedor: {numero_factura}",
            usuario_id=current_user.id
        )
        session.add(mov)
        
        # Registrar Lote automáticamente
        from datetime import date, timedelta
        lote_venc = date.today() + timedelta(days=365) # default 1 year
        lote = Lote(
            producto_id=prod_id,
            numero_lote=f"LOT-FAC-{numero_factura}",
            fecha_ingreso=date.today(),
            fecha_vencimiento=lote_venc,
            stock_inicial=qty,
            stock_actual=qty,
            costo_unitario=costo
        )
        session.add(lote)
        
        # Registrar en Historial de Costos del Producto si varió
        prod = session.get(Producto, prod_id)
        if prod and prod.precio_costo != costo:
            hist_c = HistorialCostos(
                producto_id=prod_id,
                proveedor_id=proveedor_id,
                costo_anterior=prod.precio_costo,
                costo_nuevo=costo,
                usuario_id=current_user.id,
                fecha_registro=get_cr_time(),
                motivo=f"Actualización por registro de compra Factura {numero_factura}"
            )
            session.add(hist_c)
            prod.precio_costo = costo
            session.add(prod)
            
    session.commit()
    
    # Auditoría
    audit = Auditoria(
        usuario_id=current_user.id,
        accion="REGISTRO_COMPRA",
        modulo="Compras",
        detalles=f"Compra registrada. Factura: {numero_factura}. Total: ₡{total}",
        ip_address=request.client.host
    )
    session.add(audit)
    session.commit()
    
    return {"message": "Compra registrada con éxito y existencias actualizadas", "id": compra.id}

@router.post("/{id}/annul", status_code=200)
def annul_purchase(
    id: int,
    request: Request,
    current_user: Usuario = Depends(PermissionChecker("purchases:access")),
    session: Session = Depends(get_session)
):
    compra = session.get(Compra, id)
    if not compra:
        raise HTTPException(status_code=404, detail="Compra no encontrada.")
        
    if compra.estado == "anulado":
        raise HTTPException(status_code=400, detail="Esta compra ya ha sido anulada.")
        
    # Verificar existencias suficientes para anular
    for d in compra.detalles:
        inv = session.exec(select(Inventario).where(
            Inventario.producto_id == d.producto_id,
            Inventario.sucursal_id == compra.sucursal_id
        )).first()
        if not inv or inv.existencia_actual < d.cantidad:
            prod_name = d.producto.nombre if d.producto else f"ID {d.producto_id}"
            raise HTTPException(
                status_code=400,
                detail=f"Stock insuficiente del producto {prod_name} para anular la compra. Disponible: {inv.existencia_actual if inv else 0}, Requerido: {d.cantidad}"
            )
            
    # Restar stock e ingresar movimientos kardex y lotes
    for d in compra.detalles:
        inv = session.exec(select(Inventario).where(
            Inventario.producto_id == d.producto_id,
            Inventario.sucursal_id == compra.sucursal_id
        )).first()
        
        inv.existencia_actual -= d.cantidad
        session.add(inv)
        
        # Kardex
        mov = MovimientoInventario(
            inventario_id=inv.id,
            tipo_movimiento="salida",
            cantidad=d.cantidad,
            motivo=f"Salida por anulación de compra. Factura Proveedor: {compra.numero_factura}",
            usuario_id=current_user.id
        )
        session.add(mov)
        
        # Lote
        lote = session.exec(select(Lote).where(
            Lote.producto_id == d.producto_id,
            Lote.numero_lote == f"LOT-FAC-{compra.numero_factura}"
        )).first()
        if lote:
            lote.stock_actual = max(0.0, lote.stock_actual - d.cantidad)
            session.add(lote)
            
    # Anular Cuenta por Pagar si existiera
    cxp = session.exec(select(CuentaPorPagar).where(CuentaPorPagar.compra_id == compra.id)).first()
    if cxp:
        cxp.estado = "anulado"
        cxp.saldo_pendiente = 0.0
        session.add(cxp)
        
    compra.estado = "anulado"
    session.add(compra)
    
    # Auditoría
    audit = Auditoria(
        usuario_id=current_user.id,
        accion="ANULAR_COMPRA",
        modulo="Compras",
        detalles=f"Compra de factura {compra.numero_factura} anulada por usuario {current_user.nombre}",
        ip_address=request.client.host
    )
    session.add(audit)
    session.commit()
    
    return {"message": "Compra anulada con éxito y existencias actualizadas"}

@router.get("/accounts-payable", response_model=List[Dict[str, Any]])
def get_accounts_payable(
    proveedor_id: Optional[int] = Query(None),
    estado: Optional[str] = Query(None),
    search: Optional[str] = Query(None, description="Buscar por número de factura o nombre de proveedor"),
    fecha_inicio: Optional[date] = Query(None),
    fecha_fin: Optional[date] = Query(None),
    monto_min: Optional[float] = Query(None),
    monto_max: Optional[float] = Query(None),
    current_user: Usuario = Depends(get_current_user),
    session: Session = Depends(get_session)
):
    query = select(CuentaPorPagar)
    if proveedor_id:
        query = query.where(CuentaPorPagar.proveedor_id == proveedor_id)
    if estado:
        query = query.where(CuentaPorPagar.estado == estado)
    if fecha_inicio:
        query = query.where(CuentaPorPagar.fecha_vencimiento >= fecha_inicio)
    if fecha_fin:
        query = query.where(CuentaPorPagar.fecha_vencimiento <= fecha_fin)
    if monto_min is not None:
        query = query.where(CuentaPorPagar.monto_total >= monto_min)
    if monto_max is not None:
        query = query.where(CuentaPorPagar.monto_total <= monto_max)
        
    cxp_list = session.exec(query).all()
    results = []
    for c in cxp_list:
        p = c.proveedor
        comp = c.compra
        if search:
            search_lower = search.lower()
            if (search_lower not in p.nombre.lower() and 
                search_lower not in p.identificacion.lower() and 
                search_lower not in comp.numero_factura.lower()):
                continue
        results.append({
            "id": c.id,
            "proveedor_id": c.proveedor_id,
            "proveedor_nombre": p.nombre,
            "compra_id": c.compra_id,
            "compra_factura": comp.numero_factura,
            "monto_total": c.monto_total,
            "saldo_pendiente": c.saldo_pendiente,
            "fecha_vencimiento": c.fecha_vencimiento,
            "estado": c.estado
        })
    return results

@router.post("/accounts-payable/{id}/pay", status_code=200)
def pay_account_payable(
    id: int,
    payload: Dict[str, Any],
    request: Request,
    current_user: Usuario = Depends(PermissionChecker("purchases:access")),
    session: Session = Depends(get_session)
):
    cxp = session.get(CuentaPorPagar, id)
    if not cxp:
        raise HTTPException(status_code=404, detail="Cuenta por pagar no encontrada")
        
    monto_pago = float(payload.get("monto", 0))
    metodo_pago = payload.get("metodo_pago", "efectivo")
    
    if monto_pago <= 0:
        raise HTTPException(status_code=400, detail="El monto del pago debe ser mayor a cero")
        
    if cxp.saldo_pendiente < monto_pago:
        raise HTTPException(status_code=400, detail=f"El monto del pago (₡{monto_pago:.2f}) excede el saldo pendiente (₡{cxp.saldo_pendiente:.2f})")
        
    cxp.saldo_pendiente -= monto_pago
    if cxp.saldo_pendiente <= 0.01:
        cxp.saldo_pendiente = 0.0
        cxp.estado = "pagada"
        
    session.add(cxp)
    
    # Si la compra está totalmente pagada, cambiar su estado
    compra = session.get(Compra, cxp.compra_id)
    if compra and cxp.estado == "pagada":
        # Check if there are other unpaid accounts payable for this compra (should be none)
        unpaid = session.exec(select(CuentaPorPagar).where(
            CuentaPorPagar.compra_id == compra.id,
            CuentaPorPagar.estado == "pendiente"
        )).all()
        # Excluir la actual que ya se modificó en memoria pero podría seguir en la base de datos
        unpaid = [u for u in unpaid if u.id != cxp.id]
        if not unpaid:
            compra.estado = "pagada"
            session.add(compra)
            
    # Registrar el egreso en el arqueo de caja si hay uno abierto
    from ..database.schema import Arqueo, DetalleArqueo
    arqueo = session.exec(select(Arqueo).where(
        Arqueo.usuario_id == current_user.id,
        Arqueo.estado == "abierta"
    )).first()
    if arqueo:
        detalle_caja = DetalleArqueo(
            arqueo_id=arqueo.id,
            tipo_movimiento=f"pago_proveedor_{metodo_pago}",
            descripcion=f"Pago a proveedor: {cxp.proveedor.nombre}. Compra Factura: {compra.numero_factura if compra else 'N/A'}",
            monto=monto_pago,
            usuario_id=current_user.id
        )
        session.add(detalle_caja)
        
    # Registrar auditoria
    audit = Auditoria(
        usuario_id=current_user.id,
        accion="PAGO_PROVEEDOR",
        modulo="Compras",
        detalles=f"Pago de cuenta por pagar ID {cxp.id} a proveedor {cxp.proveedor.nombre} por ₡{monto_pago:.2f} via {metodo_pago}. Saldo pendiente: ₡{cxp.saldo_pendiente:.2f}",
        ip_address=request.client.host
    )
    session.add(audit)
    
    # Enviar correo electrónico real para el abono de compras
    correo_destino = payload.get("correo_destino") or (cxp.proveedor.correo if (cxp.proveedor and cxp.proveedor.correo) else None)
    
    enviar_correo = payload.get("enviar_correo")
    if enviar_correo is None:
        enviar_correo = bool(correo_destino)
    elif isinstance(enviar_correo, str):
        enviar_correo = enviar_correo.lower() == "true"
    else:
        enviar_correo = bool(enviar_correo)
        
    if enviar_correo and correo_destino:
        subject = f"Comprobante de pago realizado - Minisúper M Y M"
        body_text = f"Estimados {cxp.proveedor.nombre},\n\nLe confirmamos que hemos registrado un pago por un monto de ₡{monto_pago:.2f} para la factura de compra #{compra.numero_factura if compra else 'N/A'}.\n\nEl saldo pendiente restante es de: ₡{cxp.saldo_pendiente:.2f}.\n\n¡Gracias por su servicio!"
        
        sent = send_email_with_pdf(correo_destino, subject, body_text, None, session)
        
        from ..database.schema import Configuracion
        smtp_host = session.exec(select(Configuracion).where(Configuracion.clave == "smtp_host")).first()
        accion_email = "ENVIO_ABONO_EMAIL_EXITOSO" if sent else "ENVIO_ABONO_EMAIL_SIMULADO"
        detalles_email = f"Envío real de comprobante de pago de cuenta por pagar (₡{monto_pago:.2f}) a {correo_destino} usando SMTP {smtp_host.valor if smtp_host else 'default'}" if sent else f"Envío simulado/fallido de comprobante de pago de cuenta por pagar (₡{monto_pago:.2f}) a {correo_destino} usando SMTP {smtp_host.valor if smtp_host else 'default'}"
        
        audit_email = Auditoria(
            usuario_id=current_user.id,
            accion=accion_email,
            modulo="Compras",
            detalles=detalles_email,
            ip_address=request.client.host
        )
        session.add(audit_email)
        
    session.commit()
    
    return {"message": "Pago a proveedor registrado con éxito", "saldo_pendiente": cxp.saldo_pendiente, "estado": cxp.estado}

@router.get("", response_model=List[Dict[str, Any]])
def get_purchases(
    search: Optional[str] = Query(None, description="Buscar por número de factura o nombre de proveedor"),
    fecha_inicio: Optional[date] = Query(None),
    fecha_fin: Optional[date] = Query(None),
    monto_min: Optional[float] = Query(None),
    monto_max: Optional[float] = Query(None),
    estado: Optional[str] = Query(None),
    current_user: Usuario = Depends(get_current_user),
    session: Session = Depends(get_session)
):
    query = select(Compra)
    if estado:
        query = query.where(Compra.estado == estado)
    if fecha_inicio:
        query = query.where(Compra.fecha_compra >= datetime.combine(fecha_inicio, datetime.min.time()))
    if fecha_fin:
        query = query.where(Compra.fecha_compra <= datetime.combine(fecha_fin, datetime.max.time()))
    if monto_min is not None:
        query = query.where(Compra.total >= monto_min)
    if monto_max is not None:
        query = query.where(Compra.total <= monto_max)
        
    compra_list = session.exec(query).all()
    results = []
    for c in compra_list:
        p = c.proveedor
        if search:
            search_lower = search.lower()
            if (search_lower not in c.numero_factura.lower() and 
                search_lower not in p.nombre.lower() and 
                search_lower not in p.identificacion.lower()):
                continue
        results.append({
            "id": c.id,
            "proveedor_id": c.proveedor_id,
            "proveedor_nombre": p.nombre,
            "numero_factura": c.numero_factura,
            "fecha_compra": c.fecha_compra,
            "total": c.total,
            "estado": c.estado
        })
    return results

@router.put("/providers/{id}", status_code=200)
def update_provider(
    id: int,
    payload: Dict[str, Any],
    request: Request,
    current_user: Usuario = Depends(PermissionChecker("purchases:access")),
    session: Session = Depends(get_session)
):
    prov = session.get(Proveedor, id)
    if not prov:
        raise HTTPException(status_code=404, detail="Proveedor no encontrado")
        
    prov.nombre = payload.get("nombre", prov.nombre)
    prov.identificacion = payload.get("identificacion", prov.identificacion)
    prov.contacto = payload.get("contacto", prov.contacto)
    prov.telefono = payload.get("telefono", prov.telefono)
    prov.correo = payload.get("correo", prov.correo)
    prov.direccion = payload.get("direccion", prov.direccion)
    prov.activo = payload.get("activo", prov.activo)
    
    session.add(prov)
    
    # Registrar auditoría
    audit = Auditoria(
        usuario_id=current_user.id,
        accion="MODIFICAR_PROVEEDOR",
        modulo="Compras",
        detalles=f"Proveedor modificado: {prov.nombre} ({prov.identificacion}).",
        ip_address=request.client.host
    )
    session.add(audit)
    session.commit()
    
    return {"message": "Proveedor actualizado con éxito"}

@router.delete("/providers/{id}", status_code=200)
def delete_provider(
    id: int,
    request: Request,
    current_user: Usuario = Depends(PermissionChecker("purchases:access")),
    session: Session = Depends(get_session)
):
    prov = session.get(Proveedor, id)
    if not prov:
        raise HTTPException(status_code=404, detail="Proveedor no encontrado")
        
    prov.activo = False
    session.add(prov)
    
    # Registrar auditoría
    audit = Auditoria(
        usuario_id=current_user.id,
        accion="ELIMINAR_PROVEEDOR",
        modulo="Compras",
        detalles=f"Proveedor desactivado (borrado lógico): {prov.nombre} ({prov.identificacion})",
        ip_address=request.client.host
    )
    session.add(audit)
    session.commit()
    
    return {"message": "Proveedor desactivado con éxito"}


def parse_text_invoice(text: str, res: Dict[str, Any]):
    import re
    # 1. Buscar número de factura
    invoice_patterns = [
        r"(?:factura|invoice|fac|consecutivo|n[úo]mero|no\.?)\s*(?:n[úo]mero|no\.?|n[ºo])?\s*[:#-]?\s*([a-zA-Z0-9-]+)",
        r"nº\s*([a-zA-Z0-9-]+)"
    ]
    for pattern in invoice_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            # Si el valor capturado es trivial como 'no' o 'numero', omitirlo y seguir buscando
            val = match.group(1).strip()
            if val.lower() not in ["no", "numero", "num"]:
                res["numero_factura"] = val
                break
            
    # 2. Buscar datos de proveedor
    provider_patterns = [
        r"([A-Za-z0-9\s,&.-]+(?:s\.a\.|sa|s\.r\.l\.|srl|ltda|limitada))",
        r"proveedor\s*:\s*([A-Za-z0-9\s,&.-]+)"
    ]
    for pattern in provider_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            res["proveedor_nombre"] = match.group(1).strip()
            break
            
    # Buscar identificación jurídica
    id_patterns = [
        r"c[ée]dula\s*jur[íi]dica\s*[:#-]?\s*([0-9-]+)",
        r"(?:id|ruc|nit|nif)\s*[:#-]?\s*([0-9-]+)"
    ]
    for pattern in id_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            res["proveedor_identificacion"] = match.group(1).strip()
            break

    # 3. Buscar líneas de productos
    lines = text.split("\n")
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        parts = re.split(r'\s*\|\s*|\s*;\s*|\t+', line_clean)
        if len(parts) < 3:
            match = re.match(r'^(\d+(?:\.\d+)?)\s+(.+?)\s+(\d+(?:\.\d+)?)(?:\s+(\d+(?:\.\d+)?))?$', line_clean)
            if match:
                try:
                    qty = float(match.group(1))
                    name = match.group(2).strip()
                    cost = float(match.group(3))
                    if len(name) > 3 and qty > 0 and cost > 0:
                        res["items"].append({
                            "producto_nombre": name,
                            "cantidad": qty,
                            "costo_unitario": cost,
                            "sku": "",
                            "codigo_barras": ""
                        })
                except:
                    pass
        else:
            candidate_qty = None
            candidate_name = None
            candidate_cost = None
            
            for part in parts:
                part_clean = part.strip()
                if not part_clean:
                    continue
                if re.match(r'^\d+(?:\.\d+)?$', part_clean) or re.match(r'^[₡$]?\s*\d+(?:\.\d+)?$', part_clean):
                    try:
                        val = float(re.sub(r'[₡$,\s]', '', part_clean))
                        if candidate_qty is None and val < 10000:
                            candidate_qty = val
                        elif candidate_cost is None:
                            candidate_cost = val
                    except:
                        pass
                else:
                    if len(part_clean) > 2:
                        candidate_name = part_clean
            
            if candidate_name and candidate_qty and candidate_cost:
                res["items"].append({
                    "producto_nombre": candidate_name,
                    "cantidad": candidate_qty,
                    "costo_unitario": candidate_cost,
                    "sku": "",
                    "codigo_barras": ""
                })
                
    res["parsed_successfully"] = len(res["items"]) > 0


@router.post("/parse-invoice", status_code=200)
def parse_invoice(
    payload: Dict[str, Any],
    request: Request,
    current_user: Usuario = Depends(PermissionChecker("purchases:access")),
    session: Session = Depends(get_session)
):
    import base64
    import io
    import os
    import xml.etree.ElementTree as ET
    import re
    
    file_base64 = payload.get("file_base64")
    file_name = payload.get("file_name", "")
    
    if not file_base64:
        raise HTTPException(status_code=400, detail="No se proporcionó el contenido del archivo.")
        
    try:
        file_bytes = base64.b64decode(file_base64)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error al decodificar base64: {str(e)}")
        
    res = {
        "numero_factura": "",
        "proveedor_nombre": "",
        "proveedor_identificacion": "",
        "items": [],
        "parsed_successfully": False,
        "method_used": ""
    }
    
    ext = os.path.splitext(file_name.lower())[1]
    
    if ext == ".xml":
        try:
            xml_str = file_bytes.decode('utf-8', errors='ignore')
            xml_clean = re.sub(r' xmlns="[^"]+"', '', xml_str)
            xml_clean = re.sub(r' xmlns:[^=]+="[^"]+"', '', xml_clean)
            
            root = ET.fromstring(xml_clean.encode('utf-8'))
            
            consecutivo_el = root.find(".//NumeroConsecutivo")
            if consecutivo_el is not None:
                res["numero_factura"] = consecutivo_el.text
            else:
                clave_el = root.find(".//Clave")
                if clave_el is not None:
                    res["numero_factura"] = clave_el.text[-10:]
                    
            emisor_el = root.find(".//Emisor")
            if emisor_el is not None:
                nombre_el = emisor_el.find(".//Nombre")
                if nombre_el is not None:
                    res["proveedor_nombre"] = nombre_el.text
                id_el = emisor_el.find(".//Identificacion/Numero")
                if id_el is not None:
                    res["proveedor_identificacion"] = id_el.text
                    
            lineas = root.findall(".//LineaDetalle")
            for linea in lineas:
                detalle_el = linea.find(".//Detalle")
                cant_el = linea.find(".//Cantidad")
                cost_el = linea.find(".//PrecioUnitario")
                
                if detalle_el is not None and cant_el is not None and cost_el is not None:
                    res["items"].append({
                        "producto_nombre": detalle_el.text,
                        "cantidad": float(cant_el.text),
                        "costo_unitario": float(cost_el.text),
                        "sku": "",
                        "codigo_barras": ""
                    })
                    
            res["parsed_successfully"] = len(res["items"]) > 0
            res["method_used"] = "XML (Factura Electrónica CR)"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error al procesar XML: {str(e)}")
            
    elif ext == ".pdf":
        try:
            import pypdf
            pdf_file = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(pdf_file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
                
            parse_text_invoice(text, res)
            res["method_used"] = "PDF Text Extraction"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error al procesar PDF: {str(e)}")
            
    elif ext in [".docx", ".doc"]:
        try:
            import docx
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            
            text_lines = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_lines.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    row_txt = " | ".join([cell.text.strip() for cell in row.cells])
                    text_lines.append(row_txt)
                    
            text = "\n".join(text_lines)
            parse_text_invoice(text, res)
            res["method_used"] = "DOCX Text Extraction"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error al procesar Word: {str(e)}")
            
    elif ext in [".png", ".jpg", ".jpeg", ".bmp"]:
        try:
            res["method_used"] = "Simulated OCR Engine"
            file_lower = file_name.lower()
            if "coca" in file_lower or "bebidas" in file_lower:
                res["proveedor_nombre"] = "FEMSA Coca-Cola"
                res["proveedor_identificacion"] = "3-101-000456"
                res["numero_factura"] = "FAC-FEMSA-8827"
                res["items"] = [
                    {"producto_nombre": "Refresco Coca-Cola 600ml", "cantidad": 24.0, "costo_unitario": 920.0, "sku": "B-COCA-600-01", "codigo_barras": "7441003502941"}
                ]
            elif "dos pinos" in file_lower or "lacteos" in file_lower or "leche" in file_lower:
                res["proveedor_nombre"] = "Cooperativa Dos Pinos"
                res["proveedor_identificacion"] = "3-104-123456"
                res["numero_factura"] = "FAC-DP-55102"
                res["items"] = [
                    {"producto_nombre": "Leche Semidescremada 1L", "cantidad": 12.0, "costo_unitario": 750.0, "sku": "L-LECHE-DES-01", "codigo_barras": "7441001157190"}
                ]
            elif "sardimar" in file_lower or "atun" in file_lower or "conservas" in file_lower:
                res["proveedor_nombre"] = "Alimentos de Centroamérica S.A."
                res["proveedor_identificacion"] = "3-101-998877"
                res["numero_factura"] = "FAC-SARDIMAR-4019"
                res["items"] = [
                    {"producto_nombre": "Atún Sardimar en Aceite 140g", "cantidad": 36.0, "costo_unitario": 1050.0, "sku": "C-ATUN-ACE-01", "codigo_barras": "7441001420915"}
                ]
            else:
                res["proveedor_nombre"] = "Distribuidora El Puerto S.A."
                res["proveedor_identificacion"] = "3-101-987654"
                res["numero_factura"] = "FAC-GEN-9938"
                res["items"] = [
                    {"producto_nombre": "Leche Descremada 1L", "cantidad": 20.0, "costo_unitario": 750.0, "sku": "L-LECHE-DES-01", "codigo_barras": "7441001157190"},
                    {"producto_nombre": "Refresco Coca-Cola 600ml", "cantidad": 48.0, "costo_unitario": 920.0, "sku": "B-COCA-600-01", "codigo_barras": "7441003502941"}
                ]
            res["parsed_successfully"] = True
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error en simulación de OCR: {str(e)}")
    else:
        raise HTTPException(status_code=400, detail=f"Formato de archivo '{ext}' no soportado para análisis.")
        
    return res


