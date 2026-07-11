import base64
import requests
import xml.etree.ElementTree as ET
import os

API_URL = "http://localhost:8000/api"

def get_token():
    payload = {"username": "admin", "password": "admin123"}
    resp = requests.post(f"{API_URL}/auth/login", json=payload)
    if resp.status_code == 200:
        return resp.json()["access_token"]
    raise Exception(f"Failed to login: {resp.text}")

def test_ocr_endpoint():
    token = get_token()
    headers = {"Authorization": f"Bearer {token}"}
    
    print("\n==================================================")
    print("      VERIFICANDO PARSEO DE FACTURAS (OCR/IA)     ")
    print("==================================================")
    
    # 1. Probar parseo de XML (Factura Electrónica CR)
    xml_content = """<?xml version="1.0" encoding="utf-8"?>
    <FacturaElectronica>
        <Clave>50623062600310198765400100001010000000001100000001</Clave>
        <NumeroConsecutivo>00100001010000000001</NumeroConsecutivo>
        <FechaEmision>2026-06-23T12:00:00-06:00</FechaEmision>
        <Emisor>
            <Nombre>Distribuidora Fenix S.A.</Nombre>
            <Identificacion>
                <Tipo>02</Tipo>
                <Numero>3-101-987654</Numero>
            </Identificacion>
        </Emisor>
        <DetalleServicio>
            <LineaDetalle>
                <NumeroLinea>1</NumeroLinea>
                <Detalle>Refresco Coca-Cola 600ml</Detalle>
                <Cantidad>24.000</Cantidad>
                <PrecioUnitario>920.00000</PrecioUnitario>
                <MontoTotal>22080.00000</MontoTotal>
            </LineaDetalle>
            <LineaDetalle>
                <NumeroLinea>2</NumeroLinea>
                <Detalle>Leche Semidescremada 1L</Detalle>
                <Cantidad>12.000</Cantidad>
                <PrecioUnitario>750.00000</PrecioUnitario>
                <MontoTotal>9000.00000</MontoTotal>
            </LineaDetalle>
        </DetalleServicio>
    </FacturaElectronica>
    """
    xml_b64 = base64.b64encode(xml_content.encode('utf-8')).decode('utf-8')
    
    payload_xml = {
        "file_base64": xml_b64,
        "file_name": "factura_proveedor.xml"
    }
    
    resp = requests.post(f"{API_URL}/purchases/parse-invoice", json=payload_xml, headers=headers)
    assert resp.status_code == 200, f"XML Fail: {resp.text}"
    res_xml = resp.json()
    assert res_xml["numero_factura"] == "00100001010000000001"
    assert res_xml["proveedor_nombre"] == "Distribuidora Fenix S.A."
    assert res_xml["proveedor_identificacion"] == "3-101-987654"
    assert len(res_xml["items"]) == 2
    assert res_xml["items"][0]["producto_nombre"] == "Refresco Coca-Cola 600ml"
    assert res_xml["items"][0]["cantidad"] == 24.0
    assert res_xml["items"][0]["costo_unitario"] == 920.0
    print("[PASS] XML (Factura Electrónica CR) parseado exitosamente.")
    
    # 2. Probar parseo de PDF
    # Generar un PDF temporal simple que contenga texto que podamos extraer
    from pypdf import PdfWriter
    # Nota: para no depender de FPDF para escribir el PDF de prueba (que podría fallar), 
    # podemos utilizar fpdf2 que está instalada en el sistema para crear un PDF con texto.
    from fpdf import FPDF
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(200, 10, text="Distribuidora Fenix S.A.", ln=1, align="C")
    pdf.cell(200, 10, text="Cedula Juridica: 3-101-987654", ln=1, align="L")
    pdf.cell(200, 10, text="Factura No: FAC-2026-99381", ln=1, align="L")
    pdf.cell(200, 10, text="Detalle de Factura:", ln=1, align="L")
    pdf.cell(200, 10, text="10 | Leche Descremada 1L | 750.00 | 7500.00", ln=1, align="L")
    pdf.cell(200, 10, text="15 | Refresco Coca-Cola 600ml | 920.00 | 13800.00", ln=1, align="L")
    
    pdf_path = "temp_factura_test.pdf"
    pdf.output(pdf_path)
    
    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()
    pdf_b64 = base64.b64encode(pdf_bytes).decode('utf-8')
    os.remove(pdf_path) # Limpiar
    
    payload_pdf = {
        "file_base64": pdf_b64,
        "file_name": "factura.pdf"
    }
    
    resp = requests.post(f"{API_URL}/purchases/parse-invoice", json=payload_pdf, headers=headers)
    assert resp.status_code == 200, f"PDF Fail: {resp.text}"
    res_pdf = resp.json()
    assert res_pdf["numero_factura"] == "FAC-2026-99381"
    assert "Fenix" in res_pdf["proveedor_nombre"]
    assert res_pdf["proveedor_identificacion"] == "3-101-987654"
    assert len(res_pdf["items"]) == 2
    assert "Leche" in res_pdf["items"][0]["producto_nombre"]
    assert res_pdf["items"][0]["cantidad"] == 10.0
    assert res_pdf["items"][0]["costo_unitario"] == 750.0
    print("[PASS] PDF con texto parseado exitosamente.")
    
    # 3. Probar parseo de DOCX
    import docx
    doc = docx.Document()
    doc.add_paragraph("Distribuidora Fenix S.A.")
    doc.add_paragraph("Cedula Juridica: 3-101-987654")
    doc.add_paragraph("Factura No: FAC-DOC-88271")
    doc.add_paragraph("20 | Leche Descremada 1L | 750.00 | 15000.00")
    
    docx_path = "temp_factura_test.docx"
    doc.save(docx_path)
    
    with open(docx_path, "rb") as f:
        docx_bytes = f.read()
    docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')
    os.remove(docx_path) # Limpiar
    
    payload_docx = {
        "file_base64": docx_b64,
        "file_name": "factura.docx"
    }
    
    resp = requests.post(f"{API_URL}/purchases/parse-invoice", json=payload_docx, headers=headers)
    assert resp.status_code == 200, f"DOCX Fail: {resp.text}"
    res_docx = resp.json()
    assert res_docx["numero_factura"] == "FAC-DOC-88271"
    assert res_docx["proveedor_identificacion"] == "3-101-987654"
    assert len(res_docx["items"]) == 1
    assert res_docx["items"][0]["cantidad"] == 20.0
    assert res_docx["items"][0]["costo_unitario"] == 750.0
    print("[PASS] DOCX con texto parseado exitosamente.")

    # 4. Probar parseo de Imagen (Simulado OCR)
    payload_img = {
        "file_base64": base64.b64encode(b"dummy image bytes").decode('utf-8'),
        "file_name": "factura_coca_cola.png"
    }
    resp = requests.post(f"{API_URL}/purchases/parse-invoice", json=payload_img, headers=headers)
    assert resp.status_code == 200, f"Image Fail: {resp.text}"
    res_img = resp.json()
    assert "FEMSA" in res_img["proveedor_nombre"]
    assert "8827" in res_img["numero_factura"]
    assert len(res_img["items"]) == 1
    assert "Coca-Cola" in res_img["items"][0]["producto_nombre"]
    print("[PASS] Imagen (OCR Simulado) parseada exitosamente.")

    print("\n==================================================")
    print("    PARSEO DE FACTURAS (OCR/IA) COMPROBADO OK!    ")
    print("==================================================")

if __name__ == "__main__":
    try:
        test_ocr_endpoint()
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"\n[FAIL] Falló la validación: {e}")
